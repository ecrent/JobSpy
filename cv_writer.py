"""Reconstruct a tailored .docx CV by modifying only mutable sections.

Strategy: open a fresh copy of the original, manipulate the XML to
remove old content paragraphs and insert new ones, preserving all
formatting, margins, fonts, and immutable sections (Education, Experience).
Sections are processed bottom-to-top so paragraph indices stay valid.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Spacing constants
SECTION_SPACING = Pt(7)    # between sections (before heading)
PARAGRAPH_SPACING = Pt(2)  # between paragraphs within a section

SECTION_HEADINGS_UPPER = {
    "PROFESSIONAL SUMMARY",
    "EDUCATION",
    "EXPERIENCE",
    "PROJECTS",
    "TECHNICAL SKILLS",
}


def _find_sections(doc):
    """Find section boundaries in the document."""
    sections = {}
    heading_order = []

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().upper()
        if text in SECTION_HEADINGS_UPPER:
            sections[text] = {"heading_idx": i}
            heading_order.append(text)

    for j, name in enumerate(heading_order):
        start = sections[name]["heading_idx"] + 1
        if j + 1 < len(heading_order):
            end = sections[heading_order[j + 1]]["heading_idx"] - 1
        else:
            end = len(doc.paragraphs) - 1
        sections[name]["start"] = start
        sections[name]["end"] = end

    return sections


def _make_run(text, bold=False):
    """Create a <w:r> element with optional bold formatting."""
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        rPr.append(OxmlElement("w:b"))
        r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _remove_content(doc, start_idx, end_idx):
    """Remove paragraph elements [start_idx .. end_idx] from document body."""
    body = doc.element.body
    elements = [doc.paragraphs[i]._element for i in range(start_idx, end_idx + 1)]
    for el in elements:
        body.remove(el)


def _insert_paragraph_after(ref_el, *runs):
    """Insert a new <w:p> with the given runs after ref_el. Returns the new element."""
    p = OxmlElement("w:p")
    for r in runs:
        p.append(r)
    ref_el.addnext(p)
    return p


from cv_tailor import SUMMARY_PREFIX, FIXED_PROJECTS


def _apply_spacing(doc):
    """Set compact spacing: larger gap before section headings, small gap within sections.

    Technical Skills lines get 2× the paragraph spacing used in Projects.
    """
    sections = _find_sections(doc)
    heading_indices = {info["heading_idx"] for info in sections.values()}

    # Build set of paragraph indices belonging to Technical Skills content
    tech_skills_indices = set()
    if "TECHNICAL SKILLS" in sections:
        ts = sections["TECHNICAL SKILLS"]
        tech_skills_indices = set(range(ts["start"], ts["end"] + 1))

    for i, p in enumerate(doc.paragraphs):
        pf = p.paragraph_format
        if i in heading_indices:
            pf.space_before = SECTION_SPACING
            pf.space_after = PARAGRAPH_SPACING
        elif i in tech_skills_indices:
            pf.space_before = PARAGRAPH_SPACING
            pf.space_after = Pt(0)
        else:
            pf.space_before = PARAGRAPH_SPACING
            pf.space_after = Pt(0)


def write_tailored_cv(source_path, output_path, tailored):
    """Create a new .docx with tailored mutable sections.

    Args:
        source_path: Path to the original base CV .docx
        output_path: Path for the tailored output .docx
        tailored: Dict with keys: summary_continuation, fifth_project, technical_skills
    """
    doc = Document(source_path)

    # --- TECHNICAL SKILLS (bottom-most mutable section → process first) ---
    sections = _find_sections(doc)
    if "TECHNICAL SKILLS" in sections:
        s = sections["TECHNICAL SKILLS"]
        if s["start"] <= s["end"]:
            _remove_content(doc, s["start"], s["end"])
        heading_el = doc.paragraphs[s["heading_idx"]]._element
        ref = heading_el
        for category, skills in tailored["technical_skills"].items():
            ref = _insert_paragraph_after(
                ref,
                _make_run(f"{category}: ", bold=True),
                _make_run(skills),
            )

    # --- PROJECTS (3 fixed + 1 generated) ---
    sections = _find_sections(doc)
    if "PROJECTS" in sections:
        s = sections["PROJECTS"]
        if s["start"] <= s["end"]:
            _remove_content(doc, s["start"], s["end"])
        heading_el = doc.paragraphs[s["heading_idx"]]._element
        ref = heading_el
        for project_text in FIXED_PROJECTS:
            ref = _insert_paragraph_after(ref, _make_run(project_text))
        # Fifth project from LLM
        ref = _insert_paragraph_after(ref, _make_run(tailored["fifth_project"]))

    # --- PROFESSIONAL SUMMARY (fixed prefix + generated continuation) ---
    sections = _find_sections(doc)
    if "PROFESSIONAL SUMMARY" in sections:
        s = sections["PROFESSIONAL SUMMARY"]
        if s["start"] <= s["end"]:
            _remove_content(doc, s["start"], s["end"])
        heading_el = doc.paragraphs[s["heading_idx"]]._element
        continuation = tailored["summary_continuation"]
        if continuation and not continuation.startswith((" ", ",")):
            continuation = " " + continuation
        full_summary = SUMMARY_PREFIX + continuation
        _insert_paragraph_after(
            heading_el,
            _make_run(full_summary),
        )

    _apply_spacing(doc)
    doc.save(output_path)
