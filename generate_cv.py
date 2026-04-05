import os
import re
import copy
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

from db import Job, SessionLocal

load_dotenv()

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel("gemini-2.5-flash")

BASE_CV_PATH = "talasli_cv.docx"
OUTPUT_DIR = "generated_cvs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def extract_cv_text(doc_path: str) -> str:
    """Extract all text from the base CV for sending to the AI."""
    doc = Document(doc_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def call_gemini(cv_text: str, job_title: str, company: str, job_description: str) -> dict:
    """
    Ask Gemini to tailor the CV for the specific job.
    Returns a dict with tailored sections.
    """
    prompt = f"""You are a professional CV/resume tailoring assistant. You will receive a base CV (in Turkish) and a job posting. Your task is to tailor the CV for this specific role.

IMPORTANT RULES:
- Output MUST be entirely in English
- Keep ALL real information — do NOT fabricate experience, skills, or qualifications
- Do NOT add skills or technologies the candidate doesn't have
- Maintain the same CV structure and sections

Respond in EXACTLY this format (no markdown, no extra text):

SUMMARY: <3-4 sentence professional summary tailored to this role>

PROGRAMMING_LANGUAGES: <reordered comma-separated list of all existing languages>
FRAMEWORKS: <reordered comma-separated list of all existing frameworks>
DATABASES: <reordered comma-separated list of all existing databases>
TOOLS: <reordered comma-separated list of all existing tools>

EXPERIENCE_1: <1-2 sentence tailored bullet for Materials Coordinator role>
EXPERIENCE_2: <1-2 sentence tailored bullet for Production Engineer role>

PROJECT_TITLE: <English title for the thesis project>
PROJECT_1: <2-3 sentence tailored description for thesis project>
PROJECT_2: <2-3 sentence tailored description for full-stack web app project>
PROJECT_3: <2-3 sentence tailored description for data pipeline project>

---
BASE CV:
{cv_text}

---
JOB POSTING:
Title: {job_title}
Company: {company}
Description:
{job_description[:3000]}"""

    response = model.generate_content(prompt)
    return parse_gemini_response(response.text)


def parse_gemini_response(text: str) -> dict:
    """Parse the flat key: value Gemini response into a dict."""
    result = {
        "summary": "",
        "skills": {
            "languages": "",
            "frameworks": "",
            "databases": "",
            "tools": "",
        },
        "experience_1": "",
        "experience_2": "",
        "project_title": "",
        "project_1": "",
        "project_2": "",
        "project_3": "",
    }

    key_map = {
        "SUMMARY": "summary",
        "PROGRAMMING_LANGUAGES": ("skills", "languages"),
        "FRAMEWORKS": ("skills", "frameworks"),
        "DATABASES": ("skills", "databases"),
        "TOOLS": ("skills", "tools"),
        "EXPERIENCE_1": "experience_1",
        "EXPERIENCE_2": "experience_2",
        "PROJECT_TITLE": "project_title",
        "PROJECT_1": "project_1",
        "PROJECT_2": "project_2",
        "PROJECT_3": "project_3",
    }

    for line in text.strip().split("\n"):
        line = line.strip()
        if not line or line.startswith("---"):
            continue

        for prefix, target in key_map.items():
            if line.upper().startswith(prefix + ":"):
                value = line[len(prefix) + 1:].strip()
                if isinstance(target, tuple):
                    result[target[0]][target[1]] = value
                else:
                    result[target] = value
                break

    return result


def create_tailored_cv(tailored: dict, job_title: str, company: str, output_path: str):
    """
    Copy the base CV and replace content with tailored text while preserving formatting.
    """
    doc = Document(BASE_CV_PATH)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # [3] ÖZGEÇMİŞ → SUMMARY
        if text == "ÖZGEÇMİŞ":
            _replace_run_text(para, "SUMMARY")

        # [4] Summary text
        elif i == 4 and text.startswith("Petrol ve"):
            _replace_run_text(para, tailored["summary"])

        # [5] EĞİTİM → EDUCATION
        elif text == "EĞİTİM":
            _replace_run_text(para, "EDUCATION")

        # [7] Bilgisayar Mühendisliği → Computer Engineering
        elif "Bilgisayar Mühendisliği" in text:
            _replace_run_text(para, text.replace(
                "Bilgisayar Mühendisliği Yüksek Lisans",
                "M.Sc. Computer Engineering"
            ).replace("2026", "2026"))

        # [8] Seçmeli Dersler → Elective Courses (education 1)
        elif i == 8 and text.startswith("Seçmeli Dersler"):
            _replace_run_text(para, "Elective Courses: Data Engineering, Advanced Cybersecurity and Cryptology")

        # [11] Petrol ve Doğal Gaz → Petroleum and Natural Gas Engineering
        elif "Petrol ve Doğal Gaz" in text:
            _replace_run_text(para, text.replace(
                "Petrol ve Doğal Gaz Mühendisliği Lisans",
                "B.Sc. Petroleum and Natural Gas Engineering"
            ))

        # [12] Seçmeli Dersler → Elective Courses (education 2)
        elif i == 12 and text.startswith("Seçmeli Dersler"):
            _replace_run_text(para, "Elective Courses: Mathematical Modeling of Hydrocarbon Reservoirs")

        # [13] DENEYİM → EXPERIENCE
        elif text == "DENEYİM":
            _replace_run_text(para, "EXPERIENCE")

        # [14] Materials Coordinator date
        elif "Materials Coordinator" in text and "Devam Ediyor" in text:
            _replace_run_text(para, text.replace("Devam Ediyor", "Present"))

        # [16] Experience 1 bullet
        elif i == 16 and text.startswith("Acted as the central"):
            if tailored["experience_1"]:
                _replace_run_text(para, tailored["experience_1"])

        # [19] Experience 2 bullet
        elif i == 19 and text.startswith("Managed end-to-end"):
            if tailored["experience_2"]:
                _replace_run_text(para, tailored["experience_2"])

        # [20] PROJELER → PROJECTS
        elif text == "PROJELER":
            _replace_run_text(para, "PROJECTS")

        # [21] Thesis title
        elif i == 21 and "Mikroservis" in text:
            if tailored["project_title"]:
                _replace_run_text(para, tailored["project_title"])

        # [22] Project 1 description
        elif i == 22 and text.startswith("Bulut tabanlı"):
            if tailored["project_1"]:
                _replace_run_text(para, tailored["project_1"])

        # [23] Project 2 description
        elif i == 23 and text.startswith("Kimlik doğrulama"):
            if tailored["project_2"]:
                _replace_run_text(para, tailored["project_2"])

        # [24] Project 3 description
        elif i == 24 and "Kafka" in text:
            if tailored["project_3"]:
                _replace_run_text(para, tailored["project_3"])

        # [25] TEKNİK YETKİNLİKLER → TECHNICAL SKILLS
        elif text == "TEKNİK YETKİNLİKLER":
            _replace_run_text(para, "TECHNICAL SKILLS")

        # [26-29] Skills lines
        elif text.startswith("Programlama Dilleri"):
            _replace_skills_line(para, "Programming Languages:", tailored["skills"]["languages"])
        elif text.startswith("Framework & Kütüphaneler"):
            _replace_skills_line(para, "Frameworks & Libraries:", tailored["skills"]["frameworks"])
        elif text.startswith("Veritabanı Sistemleri"):
            _replace_skills_line(para, "Database Systems:", tailored["skills"]["databases"])
        elif text.startswith("Araçlar & Platformlar"):
            _replace_skills_line(para, "Tools & Platforms:", tailored["skills"]["tools"])

    doc.save(output_path)


def _replace_run_text(para, new_text: str):
    """Replace all text in a paragraph with new text, preserving first run's formatting."""
    from lxml import etree
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    if not para.runs:
        # No runs — text is in raw XML. Clear all <w:r> child elements and add a new run.
        for r_elem in list(para._p.findall("w:r", nsmap)):
            para._p.remove(r_elem)
        run = para.add_run(new_text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return

    # Capture formatting from first run
    first_run = para.runs[0]
    font_name = first_run.font.name
    font_size = first_run.font.size
    font_bold = first_run.bold
    font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None

    # Clear ALL direct child <w:r> elements from the paragraph XML
    for r_elem in list(para._p.findall("w:r", nsmap)):
        para._p.remove(r_elem)

    # Add a fresh run with the original formatting
    new_run = para.add_run(new_text)
    if font_name:
        new_run.font.name = font_name
    if font_size:
        new_run.font.size = font_size
    if font_bold is not None:
        new_run.bold = font_bold
    if font_color:
        new_run.font.color.rgb = font_color


def _replace_skills_line(para, label: str, value: str):
    """Replace skills line preserving bold label + normal value formatting."""
    if not para.runs or not value:
        return
    # Find the first run (bold label) and set it, rest becomes value
    runs = para.runs
    if len(runs) >= 2:
        runs[0].text = label + " "
        runs[1].text = value
        for run in runs[2:]:
            run.text = ""
    else:
        runs[0].text = f"{label} {value}"


def generate_cv_for_job(job: Job) -> str:
    """Generate a tailored CV for a single job. Returns output file path."""
    cv_text = extract_cv_text(BASE_CV_PATH)

    print(f"  Calling Gemini for: {job.title} at {job.company}...")
    tailored = call_gemini(
        cv_text=cv_text,
        job_title=job.title,
        company=job.company,
        job_description=job.description or "",
    )

    # Sanitize filename
    safe_company = re.sub(r'[^\w\s-]', '', job.company or "unknown")[:30].strip()
    safe_title = re.sub(r'[^\w\s-]', '', job.title or "unknown")[:40].strip()
    filename = f"CV_{safe_company}_{safe_title}.docx".replace(" ", "_")
    output_path = os.path.join(OUTPUT_DIR, filename)

    print(f"  Creating tailored CV: {filename}")
    create_tailored_cv(tailored, job.title, job.company, output_path)

    return output_path


def main():
    session = SessionLocal()

    # Get all approved jobs (or 'new' if you haven't reviewed yet — change filter as needed)
    approved_jobs = session.query(Job).filter(Job.status == "approved").all()

    if not approved_jobs:
        print("No approved jobs found.")
        print("Go to Supabase → Table Editor → jobs → set status to 'approved' for jobs you want.")
        # Offer to process a sample job for testing
        sample = session.query(Job).filter(Job.status == "new").first()
        if sample:
            print(f"\nWould you like to test with one job? Found: '{sample.title}' at {sample.company}")
            print("Setting this job to 'approved' and generating a CV...")
            sample.status = "approved"
            session.commit()
            approved_jobs = [sample]
        else:
            session.close()
            return

    print(f"\nFound {len(approved_jobs)} approved jobs. Generating tailored CVs...\n")

    for job in approved_jobs:
        try:
            output_path = generate_cv_for_job(job)
            job.status = "cv_generated"
            session.commit()
            print(f"  Done! Saved to {output_path}\n")
        except Exception as e:
            print(f"  Error generating CV for '{job.title}': {e}\n")
            session.rollback()

    session.close()

    # Summary
    cvs = list(Path(OUTPUT_DIR).glob("*.docx"))
    print(f"{'='*60}")
    print(f"Generated {len(cvs)} tailored CVs in '{OUTPUT_DIR}/'")
    for cv in cvs:
        print(f"  - {cv.name}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
